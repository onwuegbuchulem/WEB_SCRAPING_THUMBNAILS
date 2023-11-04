import os
import random
from docx import Document
import pytesseract
from PIL import Image
import requests
import bs4
import re

# List of Wikipedia URLs to scrape (add more as needed)
wikipedia_urls = [
    "https://en.wikipedia.org/wiki/Artificial_intelligence",
    "https://en.wikipedia.org/wiki/Machine_learning",
    "https://en.wikipedia.org/wiki/Data_science",
    "https://en.wikipedia.org/wiki/Computer_programming",
    # Add more URLs as needed
]

# Function to extract sentences from a Wikipedia page
def extract_sentences_from_wikipedia(url):
    page = requests.get(url)
    
    if page.status_code == 200:
        soup = bs4.BeautifulSoup(page.text, 'html.parser')
        texts = " ".join([element.text for element in soup.find_all('p')])
        return texts
    else:
        print(f"Failed to retrieve the Wikipedia page: {url}")
        return []

# Function to remove [] from the text
def remove_numbers_and_brackets(text):
    pattern = r'\[\d+\]'
    result = re.sub(pattern, '', text)
    return result

# Function to extract sentences ending with a period
def extract_sentences(text):
    sentence_pattern = r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s'
    sentences = re.split(sentence_pattern, text)
    return sentences

# Function to randomly extract and save sentences to docx files
def random_extract_and_save(sentences, num_files, sentences_per_file, docx_folder="docx_files"):
    os.makedirs(docx_folder, exist_ok=True)

    for i in range(num_files):
        selected_sentences = random.choices(sentences, k=sentences_per_file)

        selected_sentences = [sentence for sentence in selected_sentences if sentence and sentence[0].isalpha()]

        if len(selected_sentences) < sentences_per_file:
            additional_sentences = random.choices(sentences, k=sentences_per_file - len(selected_sentences))
            selected_sentences.extend(additional_sentences)

        selected_text = '\n'.join(selected_sentences)

        doc = Document()
        doc.add_paragraph(selected_text)

        # Set the author name and make the file editable
        doc.core_properties.author = "Microsoft Office"

        docx_filename = os.path.join(docx_folder, f"{i + 1}.docx")
        doc.save(docx_filename)

    print(f"{num_files} DOCX files, each containing {sentences_per_file} random sentences (beginning with letters), saved to the DOCX folder.")

# Define the number of DOCX files to create
num_files = 10
sentences_per_file = 15

# Create a folder for DOCX files
docx_folder = "docx_files"

# Call the function random_extract_and_save for each URL
for url in wikipedia_urls:
    result = remove_numbers_and_brackets(extract_sentences_from_wikipedia(url))
    sentences = extract_sentences(result)

    # Call the function random_extract_and_save() for the current URL
    random_extract_and_save(sentences, num_files, sentences_per_file, docx_folder)

# Function to generate thumbnails for DOCX files
def generate_thumbnail(docx_folder, thumbnail_folder):
    os.makedirs(thumbnail_folder, exist_ok=True)

    for i in range(1, num_files + 1):
        docx_filename = os.path.join(docx_folder, f"{i}.docx")
        thumbnail_filename = os.path.join(thumbnail_folder, f"thumbnail_{i}.jpg")

        # Convert the first page of the DOCX to an image
        doc = Document(docx_filename)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            if len(text) >= 15:
                break
        text = '\n'.join(text)

        # Create an image from the text using pytesseract and PIL
        image = Image.new('RGB', (600, 800), (255, 255, 255))
        pytesseract.get_tesseract_version()
        image_path = "image.jpeg"
        pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
        image = pytesseract.image_to_string(image_path)
        image.save(thumbnail_filename, "JPEG")

# Define the thumbnail folder
thumbnail_folder = "thumbnails"


# Generate thumbnails for the DOCX files
generate_thumbnail(docx_folder, thumbnail_folder)
